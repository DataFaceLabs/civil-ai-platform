"""Export quality gates: address pin, site_payload mapping, derived recommendations."""

from __future__ import annotations

from io import BytesIO

import docx
import pytest
from fastapi.testclient import TestClient

from civilai_platform.app import create_app
from civilai_platform.services.export.context import (
    _MISSING,
    _PENDING_PLACEHOLDER,
    _enrich_floodplain_narration,
    _scrub_address_mash_from_narration,
    build_export_context,
    derive_recommendations,
    editor_body_to_text,
    pin_narration_to_canonical_address,
)
from civilai_platform.services.export.polish import polish_export_docx
from civilai_platform.store import get_store
from tests.conftest import bootstrap_client_user


@pytest.fixture(autouse=True)
def _dev_env(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("CIVILAI_DEV_AUTH", "true")
    monkeypatch.setenv("CIVILAI_STORE_BACKEND", "memory")
    monkeypatch.setenv("CIVILAI_ARTIFACT_BACKEND", "memory")
    monkeypatch.delenv("CIVILAI_EXPORT_ASYNC", raising=False)
    monkeypatch.delenv("CIVILAI_EXPORT_PDF_FUNCTION", raising=False)
    get_store.cache_clear()


@pytest.fixture
def client() -> TestClient:
    with TestClient(create_app()) as test_client:
        yield test_client


def test_pin_narration_rewrites_alternate_situs() -> None:
    body = (
        "The property is located at 611 E Braker Ln, Austin, Texas 78753 in Travis County. "
        "Frontage on FARLEY DR is assumed."
    )
    canonical = "612 Farley Drive, Austin, Texas 78753, United States"
    out = pin_narration_to_canonical_address(body, canonical)
    assert "611 E Braker" not in out
    assert "612 Farley Drive, Austin, Texas 78753, United States" in out
    assert "FARLEY DR" in out  # road name, not a full street address phrase


def test_derive_recommendations_never_pending() -> None:
    bullets = derive_recommendations(
        fields={"PLATTING_STATUS": "undetermined"},
        bodies={},
        property_address="612 Farley Drive, Austin, TX 78753",
    )
    assert bullets
    assert all(_PENDING_PLACEHOLDER.lower() not in b.lower() for b in bullets)
    assert any("plat" in b.lower() for b in bullets)


def test_export_context_pins_body_and_fills_site_payload(
    client: TestClient,
) -> None:
    user_id = "user-export-quality"
    bootstrap = bootstrap_client_user(
        client,
        user_id,
        email="quality@example.com",
        name="Quality Firm",
    )
    tenant_id = bootstrap["memberships"][0]["tenant_id"]
    headers = {"X-Dev-User-Id": user_id, "X-Tenant-Id": tenant_id}
    farley = "612 Farley Drive, Austin, Texas 78753, United States"
    project_id = client.post(
        "/v1/projects",
        json={
            "name": farley,
            "address": farley,
            "jurisdiction": "Travis County; City of Austin",
            "proposed_use": "Retail",
        },
        headers=headers,
    ).json()["project_id"]

    store = get_store()
    state = store.get_project_state(tenant_id, project_id)
    assert state
    sections = [
        section.model_copy(
            update={
                "body": (
                    "<p>The property is located at 611 E Braker Ln, Austin, Texas 78753 "
                    "in Travis County.</p>"
                    if section.step_key == "parcel"
                    else (
                        f"<p><em>{_PENDING_PLACEHOLDER}</em></p>"
                        if section.step_key == "exhibits"
                        else section.body
                    )
                )
            }
        )
        for section in state.sections
    ]
    state = state.model_copy(
        update={
            "sections": sections,
            "site_payload": {
                "entity_id": "entity-farley",
                "environmental": [
                    {"code": "ECOREGION", "value": "Texas Blackland Prairies"},
                    {"code": "MIN_ELEVATION", "value": "720"},
                    {"code": "MAX_ELEVATION", "value": "730"},
                    {"code": "SOIL_PRIMARY_NAME", "value": "Urban land"},
                    {"code": "SOIL_HYDROLOGIC_GROUP", "value": "D"},
                ],
            },
        }
    )
    store.put_project_state(state)

    ctx = build_export_context(
        store,
        tenant_id=tenant_id,
        project_id=project_id,
        skin_id="civil1_study_v1",
        data_api_base=None,
        job_id="job-quality",
    )
    assert ctx.template_values["property_address"] == farley
    assert "611 E Braker" not in (ctx.template_values.get("adjacent_props") or "")
    assert farley in (ctx.template_values.get("adjacent_props") or "")
    assert ctx.template_values["ecoregion"] == "Texas Blackland Prairies"
    assert ctx.template_values["min_elevation"] == "720"
    assert ctx.template_values["soil_types"] == "Urban land"
    assert ctx.template_values["recommendation_1"]
    assert _PENDING_PLACEHOLDER not in ctx.template_values["recommendation_1"]
    assert ctx.provenance.get("recommendations_derived") is True


def test_polish_scrubs_inline_missing_sentences() -> None:
    document = docx.Document()
    document.add_paragraph(
        f"The property was Travis County; City of Austin. {_MISSING}"
    )
    document.add_paragraph(_MISSING)
    document.add_paragraph(f"Elevation ranges from {_MISSING} to {_MISSING} ft.")
    before = BytesIO()
    document.save(before)

    polished = polish_export_docx(before.getvalue())
    out = docx.Document(BytesIO(polished))
    texts = [p.text.strip() for p in out.paragraphs if p.text.strip()]
    assert all(_MISSING not in t for t in texts)
    assert any("is located in Travis County" in t for t in texts)


def test_polish_clock_tower_presentation_defects() -> None:
    document = docx.Document()
    document.add_heading("3.1 Zoning", level=2)
    document.add_paragraph("Zoning")
    document.add_paragraph("Zoning District")
    document.add_paragraph(
        "According to information provided, the property is zoned Commercial Highway (CH)."
    )
    document.add_heading("3.2 Platting", level=2)
    document.add_heading("3.7 Right of Way", level=2)
    document.add_paragraph(
        "Road class: Local street (inferred)\n"
        "The road is inferred to be a local street. Frontage on Clock Tower Drive."
    )
    document.add_paragraph(
        "According to the maps and additional data the property Surface hydrology "
        "characterization pending USGS NHD overlay."
    )
    document.add_paragraph("1 Purpose Scope")
    before = BytesIO()
    document.save(before)

    polished = polish_export_docx(before.getvalue())
    out = docx.Document(BytesIO(polished))
    texts = [p.text.strip() for p in out.paragraphs if p.text.strip()]
    joined = "\n".join(texts)

    assert "Zoning District" not in texts
    assert texts.count("Zoning") == 0
    assert "3.2 Platting" not in joined  # empty leaf collapsed
    assert "Local street (inferred)" not in joined
    assert "inferred to be a local street" not in joined.lower()
    assert "Local Road" in joined
    assert "the property Surface" not in joined
    assert "Purpose & Scope" in joined
    assert "is zoned Commercial Highway" in joined


def test_polish_trappers_atx_presentation_defects() -> None:
    """Regression for Trappers Trail ATX PDF review (2026-07-23)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = docx.Document()

    # Header-only identification table (pre-fix template behavior).
    table = document.add_table(rows=2, cols=5)
    for cell, label in zip(
        table.rows[0].cells,
        ["TRACT", "PARCEL ID", "ADDRESS", "DEED DOC. NO.", "ACRES"],
        strict=True,
    ):
        cell.text = label
    # Empty data row — polish should drop the whole table.
    document.add_paragraph(
        "According to the maps and additional data the property."
    )
    document.add_paragraph(
        "The lot is currently developed. (See Exhibits: Vicinity Map, Zoning Map)."
    )
    document.add_paragraph(
        "Access fronts **Lockwood Springs**, classified as a **local road**."
    )
    document.add_paragraph("## Parcel Characteristics")
    document.add_paragraph(
        "The property exhibits modest relief across the tract."
    )
    document.add_paragraph("## Floodplain")
    document.add_heading("EXHIBITS", level=2)
    document.add_paragraph("List of Exhibits")
    empty_numbered = document.add_paragraph("")
    # Attach Word numbering so an empty para still renders as "1."
    p_pr = empty_numbered._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)

    before = BytesIO()
    document.save(before)
    polished = polish_export_docx(before.getvalue())
    out = docx.Document(BytesIO(polished))
    texts = [p.text.strip() for p in out.paragraphs if p.text.strip()]
    joined = "\n".join(texts)

    assert not out.tables  # header-only TRACT table removed
    assert "According to the maps and additional data the property." not in joined
    assert "See Exhibits" not in joined
    assert "**" not in joined
    assert "##" not in joined
    assert "Lockwood Springs" in joined
    assert "modest relief" in joined
    assert "Parcel Characteristics" not in joined
    assert "Floodplain" not in texts
    assert "List of Exhibits" not in texts
    # No orphan empty numbered paragraph left behind.
    for paragraph in out.paragraphs:
        if paragraph.text.strip():
            continue
        p_pr = paragraph._element.find(qn("w:pPr"))
        if p_pr is not None:
            assert p_pr.find(qn("w:numPr")) is None


def test_linter_allows_outline_subsequence_after_polish() -> None:
    from civilai_platform.services.export.linter import lint_docx
    from civilai_platform.services.export.skins import get_skin

    document = docx.Document()
    document.add_heading("1. Purpose and Scope", level=2)
    document.add_paragraph("Scope body.")
    document.add_heading("2. Description of the Property", level=2)
    document.add_paragraph("Property body.")
    document.add_heading("3. Feasibility Study", level=2)
    document.add_paragraph("Feasibility body.")
    document.add_heading("4. Summary", level=2)
    document.add_paragraph("Summary body.")
    document.add_heading("EXHIBITS", level=2)
    document.add_paragraph("Exhibits.")
    buf = BytesIO()
    document.save(buf)
    findings = lint_docx(buf.getvalue(), get_skin("atxcivil_v1"))
    assert "heading_sequence" not in {f["check"] for f in findings}


def test_editor_body_strips_markdown_heading_lines() -> None:
    text = editor_body_to_text(
        "## Parcel Characteristics\n\nThe tract slopes gently.\n\n## Floodplain\n\n"
        "Mapped Zone X applies."
    )
    assert "##" not in text
    assert "Parcel Characteristics" not in text
    assert "Floodplain" not in text
    assert "The tract slopes gently." in text
    assert "Mapped Zone X applies." in text


def test_scrub_address_mash_keeps_legitimate_situs_sentence() -> None:
    address = "20401 TRAPPERS TRL, Manor, TX"
    body = (
        f"The site at {address} fronts Lockwood Springs Road.\n\n"
        f"Contact Travis County Development Services, {address}, 512-854-7425 for permits."
    )
    cleaned = _scrub_address_mash_from_narration(body, address)
    assert "fronts Lockwood Springs Road" in cleaned
    assert address in cleaned  # kept in the situs sentence
    assert "Development Services" in cleaned
    # Mashed copy after the agency name should be gone.
    assert f"Development Services, {address}" not in cleaned


def test_enrich_floodplain_appends_firm_panel_once() -> None:
    prose = "The tract is mapped Zone X outside the SFHA."
    fields = {
        "panel_id": "48453C0710J",
        "effective_date": "2016-01-06T00:00:00",
    }
    once = _enrich_floodplain_narration(prose, fields)
    assert "48453C0710J" in once
    assert "2016-01-06" in once
    twice = _enrich_floodplain_narration(once, fields)
    assert twice.count("48453C0710J") == 1
