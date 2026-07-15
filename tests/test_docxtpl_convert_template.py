"""Tests for the E1 docxtpl template conversion (scripts/docxtpl_convert_template.py)."""

from __future__ import annotations

import importlib.util
import sys
import tempfile
from pathlib import Path

import docx
import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]


def _load_convert_module():
    module_path = REPO_ROOT / "scripts" / "docxtpl_convert_template.py"
    spec = importlib.util.spec_from_file_location("docxtpl_convert_template", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError("unable to load docxtpl_convert_template module")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _write_source(path: Path) -> None:
    document = docx.Document()
    # A token alone in its own paragraph, in SUBDOC_TOKENS -- must become a
    # paragraph tag `{{p ... }}`.
    document.add_paragraph("{WATER_SERVICE}")
    # A token inline with other text, not in SUBDOC_TOKENS -- must become an
    # inline tag `{{ ... }}`.
    document.add_paragraph("The property consists of {PROPERTY_ACRES} acres.")
    # A token alone in its own paragraph, but NOT in SUBDOC_TOKENS -- still an
    # inline tag, since only SUBDOC_TOKENS get the paragraph-tag treatment.
    document.add_paragraph("{PLATTING_STATUS}")
    document.save(str(path))


def test_subdoc_token_alone_in_paragraph_gets_paragraph_tag(tmp_path: Path) -> None:
    mod = _load_convert_module()
    source = tmp_path / "source.docx"
    dest = tmp_path / "dest.docx"
    _write_source(source)

    converted, residue = mod.convert(source, dest)

    assert converted == 3
    assert residue == []
    saved = docx.Document(str(dest))
    assert saved.paragraphs[0].text.strip() == "{{p water_service }}"


def test_non_subdoc_token_gets_inline_tag_even_when_alone(tmp_path: Path) -> None:
    mod = _load_convert_module()
    source = tmp_path / "source.docx"
    dest = tmp_path / "dest.docx"
    _write_source(source)

    mod.convert(source, dest)

    saved = docx.Document(str(dest))
    assert saved.paragraphs[2].text.strip() == "{{ platting_status }}"


def test_inline_token_with_surrounding_text_gets_inline_tag(tmp_path: Path) -> None:
    mod = _load_convert_module()
    source = tmp_path / "source.docx"
    dest = tmp_path / "dest.docx"
    _write_source(source)

    mod.convert(source, dest)

    saved = docx.Document(str(dest))
    assert saved.paragraphs[1].text == "The property consists of {{ property_acres }} acres."


def test_real_ace_template_has_no_run_split_tokens() -> None:
    # Regression guard for the E1 spike's core assumption: every {TOKEN} in the
    # real firm template sits intact within a single python-docx run, so a
    # per-run regex substitution is safe without docxtpl's run-repair dance.
    mod = _load_convert_module()
    if not mod.SOURCE_TEMPLATE.exists():
        pytest.skip("client-data/ATXCivil_Feasibility_Template.docx not present in this checkout")

    with tempfile.TemporaryDirectory() as tmp:
        dest = Path(tmp) / "converted.docx"
        converted, residue = mod.convert(mod.SOURCE_TEMPLATE, dest)
        assert converted > 0
        assert residue == []
