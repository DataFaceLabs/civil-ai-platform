"""X1 export API: real DOCX, editor narration, BYO exhibits, and status polling."""

from __future__ import annotations

import base64
from io import BytesIO

import docx
import pytest
from fastapi.testclient import TestClient

from civilai_platform.app import create_app
from civilai_platform.models.entities import MapExhibit
from civilai_platform.services.export import service as export_svc
from civilai_platform.store import get_store
from tests.conftest import bootstrap_client_user

_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


@pytest.fixture(autouse=True)
def _dev_env(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("CIVILAI_DEV_AUTH", "true")
    monkeypatch.setenv("CIVILAI_STORE_BACKEND", "memory")
    monkeypatch.setenv("CIVILAI_ARTIFACT_BACKEND", "memory")
    monkeypatch.delenv("CIVILAI_EXPORT_ASYNC", raising=False)
    get_store.cache_clear()


@pytest.fixture
def client() -> TestClient:
    with TestClient(create_app()) as test_client:
        yield test_client


def test_create_export_renders_real_docx_and_byo_exhibit(client: TestClient) -> None:
    user_id = "user-export"
    bootstrap = bootstrap_client_user(
        client,
        user_id,
        email="export@example.com",
        name="Export Firm",
    )
    tenant_id = bootstrap["memberships"][0]["tenant_id"]
    headers = {"X-Dev-User-Id": user_id, "X-Tenant-Id": tenant_id}
    created = client.post(
        "/v1/projects",
        json={
            "name": "FM 812 Feasibility",
            "address": "13903 FM 812 Rd, Del Valle, TX",
            "jurisdiction": "Travis County",
        },
        headers=headers,
    )
    assert created.status_code == 201
    project_id = created.json()["project_id"]

    store = get_store()
    state = store.get_project_state(tenant_id, project_id)
    assert state
    sections = [
        section.model_copy(
            update={
                "body": (
                    "<p><strong>Zoning:</strong> The property is outside municipal "
                    "zoning authority.</p>"
                    if section.step_key == "zoning"
                    else section.body
                )
            }
        )
        for section in state.sections
    ]
    thumbnail = f"data:image/png;base64,{base64.b64encode(_PNG).decode()}"
    state = state.model_copy(
        update={
            "sections": sections,
            "proposed_use": "commercial site development",
            "site_payload": {"entity_id": "entity-123", "serving_source": "snapshot-test"},
            "map_exhibits": [
                MapExhibit(
                    id="ex-1",
                    label="FEMA Flood Map",
                    name="fema.png",
                    size=len(_PNG),
                    mime_type="image/png",
                    thumbnail_data_url=thumbnail,
                )
            ],
        }
    )
    store.put_project_state(state)

    response = client.post(
        f"/v1/projects/{project_id}/exports",
        json={},
        headers=headers,
    )
    assert response.status_code == 201, response.text
    job = response.json()
    assert job["status"] == "succeeded"
    assert job["skin_id"] == "atxcivil_v1"
    assert job["docx_s3_key"].endswith("/study.docx")
    assert job["provenance"]["entity_id"] == "entity-123"
    assert job["provenance"]["serving_source"] == "snapshot-test"
    checks = {finding["check"] for finding in job["findings"]}
    assert "placeholder_leak" not in checks
    assert "uuid_leak" not in checks

    fetched = client.get(
        f"/v1/projects/{project_id}/exports/{job['job_id']}",
        headers=headers,
    )
    assert fetched.status_code == 200
    assert fetched.json()["docx_s3_key"] == job["docx_s3_key"]

    downloaded = client.get(
        f"/v1/projects/{project_id}/artifacts/download",
        params={"key": job["docx_s3_key"]},
        headers=headers,
    )
    assert downloaded.status_code == 200
    document = docx.Document(BytesIO(downloaded.content))
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "outside municipal zoning authority" in full_text
    assert "EXHIBIT 1 - FEMA Flood Map" in full_text
    assert any(shape for shape in document.inline_shapes)


def test_civil1_skin_export_renders_clean(client: TestClient) -> None:
    user_id = "user-civil1"
    bootstrap = bootstrap_client_user(
        client,
        user_id,
        email="civil1@example.com",
        name="Civil1 Firm",
    )
    tenant_id = bootstrap["memberships"][0]["tenant_id"]
    headers = {"X-Dev-User-Id": user_id, "X-Tenant-Id": tenant_id}
    project_id = client.post(
        "/v1/projects",
        json={
            "name": "Civil1 Skin Check",
            "address": "20401 Trappers Trail, Manor, TX",
            "jurisdiction": "Travis County",
        },
        headers=headers,
    ).json()["project_id"]

    store = get_store()
    state = store.get_project_state(tenant_id, project_id)
    assert state
    sections = [
        section.model_copy(
            update={
                "body": (
                    "<p>The site drains to a <strong>classified waterway</strong>; "
                    "compliance paths are detailed below.</p>"
                    if section.step_key == "zoning"
                    else section.body
                )
            }
        )
        for section in state.sections
    ]
    store.put_project_state(
        state.model_copy(
            update={"sections": sections, "proposed_use": "24-unit multifamily"}
        )
    )

    job = client.post(
        f"/v1/projects/{project_id}/exports",
        json={"skin_id": "civil1_study_v1"},
        headers=headers,
    ).json()
    assert job["status"] == "succeeded"
    assert job["skin_id"] == "civil1_study_v1"
    checks = {finding["check"] for finding in job["findings"]}
    assert "placeholder_leak" not in checks, job["findings"]
    assert "uuid_leak" not in checks
    assert "heading_sequence" not in checks, job["findings"]

    downloaded = client.get(
        f"/v1/projects/{project_id}/artifacts/download",
        params={"key": job["docx_s3_key"]},
        headers=headers,
    )
    document = docx.Document(BytesIO(downloaded.content))
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "classified waterway" in full_text  # narration subdoc rendered
    assert "Executive Summary" in full_text
    assert "Prepared with Civil1" not in full_text  # footer text lives in footer part
    footer_text = "\n".join(
        paragraph.text
        for section in document.sections
        for paragraph in section.footer.paragraphs
    )
    assert "Prepared with Civil1" in footer_text
    assert "{{" not in full_text and "}}" not in full_text


def test_tenant_export_skin_preference_selects_civil1(client: TestClient) -> None:
    bootstrap = bootstrap_client_user(client, "user-skin-pref", name="Pref Firm")
    tenant_id = bootstrap["memberships"][0]["tenant_id"]
    headers = {"X-Dev-User-Id": "user-skin-pref", "X-Tenant-Id": tenant_id}
    store = get_store()
    tenant = store.get_tenant(tenant_id)
    store.put_tenant(tenant.model_copy(update={"export_skin": "civil1_study_v1"}))
    project_id = client.post(
        "/v1/projects",
        json={"name": "Pref", "address": "3 Main St"},
        headers=headers,
    ).json()["project_id"]
    job = client.post(
        f"/v1/projects/{project_id}/exports", json={}, headers=headers
    ).json()
    assert job["skin_id"] == "civil1_study_v1"
    assert job["status"] == "succeeded"


def test_export_is_tenant_scoped(client: TestClient) -> None:
    first = bootstrap_client_user(client, "export-a", name="Firm A")
    second = bootstrap_client_user(client, "export-b", name="Firm B")
    tenant_a = first["memberships"][0]["tenant_id"]
    tenant_b = second["memberships"][0]["tenant_id"]
    headers_a = {"X-Dev-User-Id": "export-a", "X-Tenant-Id": tenant_a}
    headers_b = {"X-Dev-User-Id": "export-b", "X-Tenant-Id": tenant_b}
    project = client.post(
        "/v1/projects",
        json={"name": "A", "address": "1 Main St"},
        headers=headers_a,
    ).json()
    job = client.post(
        f"/v1/projects/{project['project_id']}/exports",
        json={},
        headers=headers_a,
    ).json()
    forbidden = client.get(
        f"/v1/projects/{project['project_id']}/exports/{job['job_id']}",
        headers=headers_b,
    )
    assert forbidden.status_code == 404


def test_async_export_returns_running_then_event_completes(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    bootstrap = bootstrap_client_user(client, "export-async", name="Async Firm")
    tenant_id = bootstrap["memberships"][0]["tenant_id"]
    headers = {"X-Dev-User-Id": "export-async", "X-Tenant-Id": tenant_id}
    project_id = client.post(
        "/v1/projects",
        json={"name": "Async", "address": "2 Main St"},
        headers=headers,
    ).json()["project_id"]

    enqueued: list[str] = []
    monkeypatch.setenv("CIVILAI_EXPORT_ASYNC", "true")
    monkeypatch.setattr(
        export_svc,
        "_enqueue_async_completion",
        lambda job: enqueued.append(job.job_id),
    )
    response = client.post(
        f"/v1/projects/{project_id}/exports",
        json={},
        headers=headers,
    )
    assert response.status_code == 201
    running = response.json()
    assert running["status"] == "running"
    assert enqueued == [running["job_id"]]

    completed = export_svc.complete_export_from_event(
        get_store(),
        {
            "tenant_id": tenant_id,
            "project_id": project_id,
            "job_id": running["job_id"],
        },
    )
    assert completed is not None
    assert completed.status.value == "succeeded"
    assert completed.docx_s3_key
