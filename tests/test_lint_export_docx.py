"""Tests for the E6 export-fidelity linter (scripts/lint_export_docx.py)."""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import docx

REPO_ROOT = Path(__file__).resolve().parents[1]


def _load_module():
    module_path = REPO_ROOT / "scripts" / "lint_export_docx.py"
    spec = importlib.util.spec_from_file_location("lint_export_docx", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError("unable to load lint_export_docx module")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _good_document() -> docx.document.Document:
    """A minimal document whose heading sequence matches the full ACE outline."""
    document = docx.Document()
    document.add_heading("1. Purpose and Scope", level=2)
    document.add_paragraph("This report is intended to provide preliminary engineering data.")
    document.add_heading("2. Description of the Property", level=2)
    document.add_heading("2.1 General Information", level=3)
    document.add_paragraph("The subject property consists of 1.5 acres.")
    document.add_heading("2.2 Site Characteristics", level=3)
    document.add_paragraph("The elevation ranges from 594.8 ft to 601.2 ft.")
    document.add_heading("2.3 Property Identification Number", level=3)
    document.add_paragraph("TCAD Property ID: 984219")
    document.add_heading("3. Feasibility Study", level=2)
    for n in range(1, 20):
        document.add_heading(f"3.{n} Section {n}", level=3)
        document.add_paragraph(f"Content for section 3.{n}.")
        if n == 13:
            document.add_heading("3.13.1 Required Permits", level=4)
            document.add_paragraph("Permit content.")
            document.add_heading("3.13.2 Permitting Contacts", level=4)
            document.add_paragraph("Contact content.")
    document.add_heading("4. Summary", level=2)
    document.add_paragraph("The proposed development appears to be feasible.")
    document.add_heading("EXHIBITS", level=2)
    document.add_paragraph("List of exhibits.")
    return document


def test_good_document_passes_heading_sequence() -> None:
    mod = _load_module()
    findings = mod.check_heading_sequence(_good_document())
    assert findings == []


def test_wrong_heading_sequence_flagged() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_heading("1. Purpose and Scope", level=2)
    document.add_heading("3. Feasibility Study", level=2)  # skips section 2 entirely
    findings = mod.check_heading_sequence(document)
    assert len(findings) == 1
    assert findings[0].check == "heading_sequence"


def test_placeholder_leak_detects_curly_token() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("Zoning: {ZONING_REGS}")
    document.add_paragraph("Watershed: {{ watershed_info }}")
    document.add_paragraph("![FIRM LOGO]")
    findings = mod.check_placeholder_leaks(document)
    checks = {f.check for f in findings}
    assert checks == {"placeholder_leak"}
    assert len(findings) == 3


def test_placeholder_leak_clean_document_passes() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("Zoning is single-family residential.")
    findings = mod.check_placeholder_leaks(document)
    assert findings == []


def test_unit_sanity_flags_implausible_low_elevation() -> None:
    mod = _load_module()
    document = docx.Document()
    # A real D8-class bug: raw meters (181.3) printed under a feet label.
    document.add_paragraph("The existing elevation ranges from 181.3 ft to 183.2 ft.")
    findings = mod.check_unit_sanity(document)
    assert len(findings) == 2
    assert all(f.check == "unit_sanity" for f in findings)


def test_unit_sanity_passes_plausible_austin_elevation() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("The existing elevation ranges from 594.8 ft to 601.2 ft.")
    findings = mod.check_unit_sanity(document)
    assert findings == []


def test_uuid_leak_detects_internal_id() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph(
        "Property ID 02e46844-c3be-51bd-a99d-97f922d850f8 (Bastrop County Appraisal District)"
    )
    findings = mod.check_uuid_leak(document)
    assert len(findings) == 1
    assert findings[0].check == "uuid_leak"


def test_uuid_leak_clean_document_passes() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("TCAD Property ID: 984219")
    findings = mod.check_uuid_leak(document)
    assert findings == []


def test_contradiction_flags_platting_status_conflict() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("Platting status could not be confirmed.")
    document.add_paragraph("The property is unplatted per available records.")
    findings = mod.check_contradictions(document)
    assert len(findings) == 1
    assert findings[0].check == "contradiction"


def test_contradiction_clean_document_passes() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("The property is platted.")
    findings = mod.check_contradictions(document)
    assert findings == []


def test_truncation_flags_document_not_ending_in_summary() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_heading("1. Purpose and Scope", level=2)
    document.add_paragraph("Some content.")
    document.add_heading("3.5 Utility Location & Availability", level=3)
    document.add_paragraph("• Contact AQUA WSC")
    findings = mod.check_truncation(document)
    assert len(findings) == 1
    assert findings[0].check == "truncation"


def test_truncation_flags_mid_list_ending() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_heading("4. Summary", level=2)
    document.add_paragraph("Recommendation 1:")
    findings = mod.check_truncation(document)
    assert len(findings) == 1
    assert "mid-list" in findings[0].detail


def test_truncation_passes_document_ending_in_exhibits() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_heading("4. Summary", level=2)
    document.add_paragraph("The development appears feasible.")
    document.add_heading("EXHIBITS", level=2)
    document.add_paragraph("Vicinity map.")
    findings = mod.check_truncation(document)
    assert findings == []


def test_bold_body_ratio_flags_blanket_bold() -> None:
    mod = _load_module()
    document = docx.Document()
    p = document.add_paragraph()
    run = p.add_run("This entire paragraph is bold, matching the html-to-docx path's habit.")
    run.bold = True
    findings = mod.check_bold_body_ratio(document)
    assert len(findings) == 1
    assert findings[0].check == "bold_body_ratio"


def test_bold_body_ratio_passes_normal_styled_body() -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("This is normal, unstyled body text with no manual bold runs.")
    findings = mod.check_bold_body_ratio(document)
    assert findings == []


def test_lint_aggregates_all_checks(tmp_path: Path) -> None:
    mod = _load_module()
    path = tmp_path / "good.docx"
    _good_document().save(str(path))
    findings = mod.lint(path)
    assert findings == []


def test_main_returns_nonzero_on_findings(tmp_path: Path, capsys) -> None:
    mod = _load_module()
    document = docx.Document()
    document.add_paragraph("{LEAKED_TOKEN}")
    path = tmp_path / "bad.docx"
    document.save(str(path))
    exit_code = mod.main([str(path)])
    assert exit_code == 1
    out = capsys.readouterr().out
    assert "placeholder_leak" in out


def test_main_returns_zero_on_clean_file(tmp_path: Path) -> None:
    mod = _load_module()
    path = tmp_path / "good.docx"
    _good_document().save(str(path))
    exit_code = mod.main([str(path)])
    assert exit_code == 0
