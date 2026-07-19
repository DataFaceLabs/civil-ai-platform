"""Civil1 export presentation polish + cover identity."""

from __future__ import annotations

from io import BytesIO

import docx
import pytest
from fastapi.testclient import TestClient

from civilai_platform.app import create_app
from civilai_platform.services.export.context import _MISSING, _normalize_identity
from civilai_platform.services.export.polish import polish_export_docx
from civilai_platform.store import get_store
from tests.conftest import bootstrap_client_user


@pytest.fixture(autouse=True)
def _dev_env(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("CIVILAI_DEV_AUTH", "true")
    monkeypatch.setenv("CIVILAI_STORE_BACKEND", "memory")
    monkeypatch.setenv("CIVILAI_ARTIFACT_BACKEND", "memory")
    monkeypatch.delenv("CIVILAI_EXPORT_ASYNC", raising=False)
    monkeypatch.delenv("CIVILAI_EXPORT_PDF_FUNCTION", raising=False)
    get_store.cache_clear()


@pytest.fixture
def client() -> TestClient:
    with TestClient(create_app()) as test_client:
        yield test_client


def test_normalize_identity_ignores_case_and_punctuation() -> None:
    assert _normalize_identity("7845 CLOCK TOWER DR, Austin, TX 78752") == _normalize_identity(
        "7845 Clock Tower Dr, Austin, TX 78752"
    )


def test_polish_strips_missing_paragraphs_keeps_outline() -> None:
    document = docx.Document()
    document.add_heading("3.2 Platting & Subdivision", level=2)
    document.add_paragraph(_MISSING)
    document.add_heading("3.3 Compatibility", level=2)
    document.add_paragraph("Compatibility standards apply near residential uses.")
    before = BytesIO()
    document.save(before)

    polished = polish_export_docx(before.getvalue())
    out = docx.Document(BytesIO(polished))
    texts = [p.text.strip() for p in out.paragraphs if p.text.strip()]
    assert _MISSING not in texts
    assert "3.2 Platting & Subdivision" in texts
    assert "3.3 Compatibility" in texts
    assert "Compatibility standards apply near residential uses." in texts


def test_polish_drops_empty_and_missing_table_rows() -> None:
    document = docx.Document()
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "PROPERTY"
    table.rows[0].cells[1].text = "7845 Clock Tower Dr"
    table.rows[1].cells[0].text = "EXISTING DEVELOPMENT"
    table.rows[1].cells[1].text = _MISSING
    table.rows[2].cells[0].text = "EXHIBIT 2"
    table.rows[2].cells[1].text = ""
    before = BytesIO()
    document.save(before)

    polished = polish_export_docx(before.getvalue())
    out = docx.Document(BytesIO(polished))
    assert len(out.tables) == 1
    assert len(out.tables[0].rows) == 1
    assert out.tables[0].rows[0].cells[1].text.strip() == "7845 Clock Tower Dr"


def test_civil1_cover_dedupes_address_named_project(client: TestClient) -> None:
    user_id = "user-cover-dedupe"
    bootstrap = bootstrap_client_user(
        client,
        user_id,
        email="cover@example.com",
        name="Platform",
    )
    tenant_id = bootstrap["memberships"][0]["tenant_id"]
    headers = {"X-Dev-User-Id": user_id, "X-Tenant-Id": tenant_id}
    address = "7845 CLOCK TOWER DR, Austin, TX 78752"
    project_id = client.post(
        "/v1/projects",
        json={
            "name": address,
            "address": address,
            "jurisdiction": "Travis County; City of Austin",
        },
        headers=headers,
    ).json()["project_id"]

    store = get_store()
    state = store.get_project_state(tenant_id, project_id)
    assert state
    store.put_project_state(
        state.model_copy(
            update={
                "proposed_use": "Retail",
                "sections": [
                    section.model_copy(
                        update={
                            "body": (
                                "<p>Only one recommendation.</p>"
                                if section.step_key == "exhibits"
                                else section.body
                            )
                        }
                    )
                    for section in state.sections
                ],
            }
        )
    )

    job = client.post(
        f"/v1/projects/{project_id}/exports",
        json={"skin_id": "civil1_study_v1"},
        headers=headers,
    ).json()
    assert job["status"] == "succeeded", job

    downloaded = client.get(
        f"/v1/projects/{project_id}/artifacts/download",
        params={"key": job["docx_s3_key"]},
        headers=headers,
    )
    document = docx.Document(BytesIO(downloaded.content))
    texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    cover_window = texts[:12]
    assert cover_window.count(address) == 1
    joined = "\n".join(texts)
    assert "Update fields in Word" not in joined
    assert "1 Purpose & Scope" in joined
    assert _MISSING not in texts
