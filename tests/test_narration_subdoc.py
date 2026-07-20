"""Narration export: TipTap paragraph breaks, mega-blob reflow, and bold runs."""

from __future__ import annotations

from io import BytesIO

import docx
from docxtpl import DocxTemplate

from civilai_platform.services.export.context import editor_body_to_text
from civilai_platform.services.export.render import (
    _narration_subdoc,
    _reflow_mega_chunk,
)
from civilai_platform.services.export.skins import ATXCIVIL_V1, CIVIL1_STUDY_V1


def test_editor_body_tiptap_paragraphs_become_blank_line_separated() -> None:
    html = "<p>First zoning finding.</p><p>Second access finding.</p>"
    text = editor_body_to_text(html)
    assert "First zoning finding." in text
    assert "Second access finding." in text
    assert "\n\n" in text
    assert text.split("\n\n") == ["First zoning finding.", "Second access finding."]


def test_editor_body_preserves_strong_as_markdown_bold() -> None:
    html = "<p><strong>Zoning:</strong> Outside municipal authority.</p>"
    text = editor_body_to_text(html)
    assert "**Zoning:**" in text
    assert "Outside municipal authority." in text


def test_editor_body_br_is_single_newline() -> None:
    html = "<p>Line one<br/>Line two</p>"
    text = editor_body_to_text(html)
    assert "Line one\nLine two" in text
    # Single paragraph block — no blank line between the two lines.
    assert "Line one\n\nLine two" not in text


def test_reflow_mega_chunk_splits_long_sentence_walls() -> None:
    sentences = [
        f"This is sentence number {i} about site access and drainage."
        for i in range(1, 8)
    ]
    blob = " ".join(sentences)
    assert len(blob) > 350
    parts = _reflow_mega_chunk(blob)
    assert len(parts) >= 3
    assert all(part.endswith(".") for part in parts)
    assert " ".join(parts) == blob


def test_reflow_skips_short_chunks() -> None:
    short = "Short. Still short. Ok."
    assert _reflow_mega_chunk(short) == [short]


def test_narration_subdoc_splits_tiptap_into_word_paragraphs() -> None:
    template = DocxTemplate(str(ATXCIVIL_V1.template_path))
    text = editor_body_to_text(
        "<p>Access paragraph one with enough words.</p>"
        "<p>Access paragraph two with different content.</p>"
    )
    subdoc = _narration_subdoc(template, text, body_font="Calisto MT")
    paras = [p.text.strip() for p in subdoc.paragraphs if p.text.strip()]
    assert paras == [
        "Access paragraph one with enough words.",
        "Access paragraph two with different content.",
    ]


def test_narration_subdoc_reflows_mega_blob_and_keeps_bold() -> None:
    sentences = [
        f"This is sentence number {i} covering zoning overlays and setbacks."
        for i in range(1, 8)
    ]
    blob = "**Zoning:** " + " ".join(sentences)
    template = DocxTemplate(str(ATXCIVIL_V1.template_path))
    subdoc = _narration_subdoc(template, blob, body_font="Calisto MT")
    paras = [p for p in subdoc.paragraphs if p.text.strip()]
    assert len(paras) >= 3
    first = paras[0]
    bold_runs = [run for run in first.runs if run.bold]
    assert any(run.text == "Zoning:" for run in bold_runs)
    assert first.runs[0].font.name == "Calisto MT"


def test_narration_subdoc_pins_civil1_body_font() -> None:
    template = DocxTemplate(str(CIVIL1_STUDY_V1.template_path))
    subdoc = _narration_subdoc(
        template,
        "Civil1 body paragraph for font pin check.",
        body_font="Source Sans 3",
    )
    paras = [p for p in subdoc.paragraphs if p.text.strip()]
    assert paras
    assert paras[0].runs[0].font.name == "Source Sans 3"


def test_narration_subdoc_round_trip_bytes_have_multiple_paragraphs() -> None:
    """Subdoc paragraphs survive into a minimal docx save (smoke for both skins)."""
    template = DocxTemplate(str(ATXCIVIL_V1.template_path))
    text = "Para A about the site.\n\nPara B about utilities."
    subdoc = _narration_subdoc(template, text, body_font="Calisto MT")
    # Materialize via a tiny wrapper document so we can inspect paragraph count.
    wrapper = docx.Document()
    for paragraph in subdoc.paragraphs:
        if not paragraph.text.strip():
            continue
        wp = wrapper.add_paragraph()
        for run in paragraph.runs:
            new_run = wp.add_run(run.text)
            new_run.bold = run.bold
            new_run.font.name = run.font.name
    buf = BytesIO()
    wrapper.save(buf)
    loaded = docx.Document(BytesIO(buf.getvalue()))
    texts = [p.text for p in loaded.paragraphs if p.text.strip()]
    assert texts == ["Para A about the site.", "Para B about utilities."]
