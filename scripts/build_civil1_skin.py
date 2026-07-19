#!/usr/bin/env python3
"""Build the Civil1 default skin master (assets/templates/civil1_study_v1.docx).

M1-DESIGN deliverable (CIVIL1-STUDY-FORMAT.md §5, §5.5). The design lives in code so it
is reproducible, diffable, and reviewable: rerunning this script regenerates the skin
byte-for-byte-equivalent (modulo docx zip timestamps).

Design tokens (§5.5):
- Typography: Libre Franklin (display/headings) + Source Sans 3 (body) — open-license,
  bundleable in the X2 export container. Word installs without them fall back; the PDF
  (canonical customer artifact) renders identically everywhere via container fonts.
- Color: ink #1A1D21, accent #144A66, rule gray #D5D9DE, caption gray #5A6672. Status
  colors are reserved for the constraint dashboard (X5, when the verdict enum lands in
  the contract context).
- Tables as designed objects: quiet horizontal rules only, shaded label column.
- Running footer: project · date · page X of Y · discreet "Prepared with Civil1".

Numbering is Civil1's own (§5 item 7 — skins may renumber; linters key on the per-skin
outline registered in skins.py). Boilerplate voice is originally written (§7 non-goal:
no ATX verbatim).

Usage:
    uv run python scripts/build_civil1_skin.py [--out assets/templates/civil1_study_v1.docx]
"""

from __future__ import annotations

import argparse
from pathlib import Path

import docx
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# ── Design tokens ────────────────────────────────────────────────────────────────────
HEADING_FONT = "Libre Franklin"
BODY_FONT = "Source Sans 3"

INK = RGBColor(0x1A, 0x1D, 0x21)
ACCENT = RGBColor(0x14, 0x4A, 0x66)
CAPTION = RGBColor(0x5A, 0x66, 0x72)
RULE_HEX = "D5D9DE"
BAND_HEX = "144A66"
PANEL_HEX = "F2F5F7"

MISSING = "Not available from current project data."


# ── Low-level helpers ────────────────────────────────────────────────────────────────
def _set_font(style, name: str) -> None:
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), name)


def _shade(cell: _Cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_borders(cell: _Cell, *, bottom: str | None = None, top: str | None = None) -> None:
    borders = OxmlElement("w:tcBorders")
    for edge, color in (("top", top), ("bottom", bottom)):
        element = OxmlElement(f"w:{edge}")
        if color:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:color"), color)
        else:
            element.set(qn("w:val"), "nil")
        borders.append(element)
    for edge in ("left", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    cell._tc.get_or_add_tcPr().append(borders)


def _strip_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _cell_margins(table: Table, *, top: int = 60, bottom: int = 60, left: int = 100, right: int = 100) -> None:
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    tbl_pr.append(margins)


def _field(paragraph: Paragraph, instruction: str) -> None:
    """Append a Word field (PAGE, NUMPAGES, TOC …) to a paragraph."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, end):
        run = paragraph.add_run()
        run._r.append(element)
        run.font.size = Pt(8)
        run.font.color.rgb = CAPTION
        run.font.name = BODY_FONT


# ── Styles ───────────────────────────────────────────────────────────────────────────
def define_styles(document: DocumentType) -> None:
    normal = document.styles["Normal"]
    _set_font(normal, BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.2
    pf.space_after = Pt(8)

    title = document.styles["Title"]
    _set_font(title, HEADING_FONT)
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(4)

    subtitle = document.styles["Subtitle"]
    _set_font(subtitle, HEADING_FONT)
    subtitle.font.size = Pt(13)
    subtitle.font.bold = False
    subtitle.font.color.rgb = CAPTION
    subtitle.paragraph_format.space_after = Pt(2)

    for level, (size, color, before, after) in {
        1: (16, ACCENT, 20, 8),
        2: (12.5, INK, 14, 6),
        3: (11, INK, 10, 4),
    }.items():
        heading = document.styles[f"Heading {level}"]
        _set_font(heading, HEADING_FONT)
        heading.font.size = Pt(size)
        heading.font.bold = True
        heading.font.color.rgb = color
        heading.font.italic = False
        hpf = heading.paragraph_format
        hpf.space_before = Pt(before)
        hpf.space_after = Pt(after)
        hpf.keep_with_next = True


# ── Building blocks ──────────────────────────────────────────────────────────────────
def para(document: DocumentType, text: str = "", *, style: str | None = None) -> Paragraph:
    return document.add_paragraph(text, style=style)


def heading(document: DocumentType, text: str, level: int) -> Paragraph:
    return document.add_heading(text, level=level)


def label_value_table(document: DocumentType, rows: list[tuple[str, str]]) -> Table:
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _strip_table_borders(table)
    _cell_margins(table)
    for index, (label, token) in enumerate(rows):
        label_cell, value_cell = table.rows[index].cells
        label_cell.width = Inches(2.1)
        value_cell.width = Inches(4.6)
        _shade(label_cell, PANEL_HEX)
        _cell_borders(label_cell, bottom=RULE_HEX)
        _cell_borders(value_cell, bottom=RULE_HEX)
        label_paragraph = label_cell.paragraphs[0]
        run = label_paragraph.add_run(label)
        run.font.name = HEADING_FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = CAPTION
        label_paragraph.paragraph_format.space_after = Pt(0)
        value_paragraph = value_cell.paragraphs[0]
        value_paragraph.add_run(token)
        value_paragraph.paragraph_format.space_after = Pt(0)
    return table


def narration_slot(document: DocumentType, token: str) -> None:
    """A docxtpl Subdoc slot — the paragraph tag must own its paragraph (E1 finding)."""
    document.add_paragraph(f"{{{{p {token} }}}}")


def accent_rule(document: DocumentType) -> None:
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), BAND_HEX)
    borders.append(bottom)
    p_pr.append(borders)
    paragraph.paragraph_format.space_after = Pt(2)


# ── Document assembly ────────────────────────────────────────────────────────────────
def build_cover(document: DocumentType) -> None:
    band = document.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    _strip_table_borders(band)
    _cell_margins(band, top=120, bottom=120)
    cell = band.rows[0].cells[0]
    cell.width = Inches(6.7)
    _shade(cell, BAND_HEX)
    firm = cell.paragraphs[0]
    run = firm.add_run("{{ firm_name }}")
    run.font.name = HEADING_FONT
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    firm.paragraph_format.space_after = Pt(0)

    for _ in range(4):
        para(document)
    para(document, "FEASIBILITY STUDY", style="Title")
    para(document, "{{ project_name }}", style="Subtitle")
    para(document, "{{ property_address }}", style="Subtitle")
    accent_rule(document)
    para(document)

    label_value_table(
        document,
        [
            ("PREPARED FOR", "{{ client_name }}"),
            ("PROPOSED DEVELOPMENT", "{{ proposed_development }}"),
            ("GOVERNING JURISDICTION", "{{ governing_juris }}"),
            ("REPORT DATE", "{{ report_date }}"),
        ],
    )

    for _ in range(6):
        para(document)
    firm_block = para(document, "{{ firm_name }}")
    firm_block.runs[0].font.bold = True
    para(document, "{{ firm_address }} · {{ firm_location }}").runs[0].font.color.rgb = CAPTION
    para(document, "{{ firm_phone }}").runs[0].font.color.rgb = CAPTION
    document.add_page_break()


def build_exec_summary(document: DocumentType) -> None:
    heading(document, "Executive Summary", 1)
    para(
        document,
        "This study screens the subject property for land-development feasibility against "
        "governed parcel, entitlement, constraint, and utility data. Findings below are "
        "summarized from the sections of this report; each fact carries its source and "
        "data vintage.",
    )
    heading(document, "Key Facts", 2)
    label_value_table(
        document,
        [
            ("PROPERTY", "{{ property_address }}"),
            ("ACREAGE", "{{ property_acres }}"),
            ("JURISDICTION", "{{ governing_juris }}"),
            ("PROPOSED USE", "{{ proposed_development }}"),
            ("EXISTING DEVELOPMENT", "{{ existing_development }}"),
            ("REPORT DATE", "{{ report_date }}"),
        ],
    )
    heading(document, "Priority Recommendations", 2)
    for index in range(1, 6):
        item = document.add_paragraph(style="List Number")
        item.add_run(f"{{{{ recommendation_{index} }}}}")
    document.add_page_break()


def build_toc(document: DocumentType) -> None:
    heading(document, "Contents", 1)
    toc = document.add_paragraph()
    _field(toc, 'TOC \\o "1-2" \\h \\z \\u')
    caption = para(document, "Update fields in Word (Ctrl+A, F9) to refresh the table of contents.")
    caption.runs[0].font.size = Pt(8)
    caption.runs[0].font.color.rgb = CAPTION
    document.add_page_break()


def build_body(document: DocumentType) -> None:
    # 1 — Purpose & Scope (originally written voice; format §7 forbids ATX verbatim).
    heading(document, "1 Purpose & Scope", 1)
    para(
        document,
        "{{ firm_name }} prepared this feasibility study to evaluate the suitability of "
        "{{ property_address }} for the proposed {{ proposed_development }}. The study "
        "compiles governed parcel, entitlement, site-constraint, and utility records; "
        "identifies conditions that materially affect development; and recommends the "
        "confirmations, studies, and permits that should precede design.",
    )
    para(
        document,
        "Facts in this report are drawn from the public and commercial sources listed with "
        "each section, as of the report date. Where a record could not be confirmed, the "
        "report says so explicitly rather than inferring a value.",
    )

    heading(document, "2 Property", 1)
    heading(document, "2.1 Site Overview", 2)
    label_value_table(
        document,
        [
            ("ACREAGE", "{{ property_acres }}"),
            ("EXISTING DEVELOPMENT", "{{ existing_development }}"),
        ],
    )
    heading(document, "2.2 Property Identification", 2)
    para(document, "{{ tcad_info }}")
    para(document, "{{ tcad_discrepancies }}")
    heading(document, "2.3 Adjacent Sites & Context", 2)
    para(document, "{{ adjacent_props }}")

    heading(document, "3 Entitlements & Administration", 1)
    heading(document, "3.1 Zoning & Land Use", 2)
    narration_slot(document, "zoning_regs")
    heading(document, "3.2 Platting & Subdivision", 2)
    para(document, "{{ platting_status }}")
    heading(document, "3.3 Compatibility & Design Standards", 2)
    para(document, "{{ compatibility_stds }}")
    heading(document, "3.4 Governing Jurisdictions", 2)
    para(document, "{{ governing_juris }}")
    para(document, "{{ jurisdiction_info }}")
    heading(document, "3.5 Required Permits", 2)
    para(document, "{{ required_permits }}")
    heading(document, "3.6 Permitting Contacts", 2)
    para(document, "{{ permit_contacts }}")
    heading(document, "3.7 Development Agreements", 2)
    para(document, "{{ dev_agreements }}")
    heading(document, "3.8 Easements & Setbacks", 2)
    para(document, "{{ easements_setbacks }}")
    heading(document, "3.9 Surveys, Title & Other Documents", 2)
    para(document, "{{ completed_docs }}")

    heading(document, "4 Site Constraints", 1)
    heading(document, "4.1 Watershed & Waterways", 2)
    para(document, "{{ watershed_info }}")
    label_value_table(
        document,
        [
            ("WATERWAY SETBACK", "{{ waterway_setback }}"),
            ("DRAINAGE AREAS", "{{ drainage_areas }}"),
        ],
    )
    heading(document, "4.2 Impervious Cover", 2)
    para(document, "{{ impervious_regs }}")
    heading(document, "4.3 Soils, Elevation & Topography", 2)
    label_value_table(
        document,
        [
            ("SOIL TYPES", "{{ soil_types }}"),
            ("HYDROLOGIC GROUP", "{{ soil_class }}"),
            ("ELEVATION RANGE", "{{ min_elevation }} – {{ max_elevation }} ft"),
            ("SLOPE RANGE", "{{ min_slope }} – {{ max_slope }} %"),
            ("ECOREGION", "{{ ecoregion }}"),
        ],
    )
    heading(document, "4.4 Floodplain Status", 2)
    narration_slot(document, "floodplain_status")
    para(document, "{{ floodplain_reqs }}")
    heading(document, "4.5 Drainage Areas & Design Criteria", 2)
    para(document, "{{ drainage_criteria }}")
    heading(document, "4.6 Water Quality & Detention", 2)
    para(document, "{{ water_quality_reqs }}")
    heading(document, "4.7 Environmental Overlays", 2)
    para(document, "{{ ecoregion_desc }}")
    label_value_table(
        document,
        [
            ("EROSION HAZARD", "{{ erosion_hazard }}"),
            ("HYDROLOGY", "{{ hydrology_char }}"),
        ],
    )

    heading(document, "5 Utilities, Access & Mobility", 1)
    heading(document, "5.1 Water Service", 2)
    narration_slot(document, "water_service")
    heading(document, "5.2 Wastewater Service", 2)
    narration_slot(document, "wastewater_service")
    heading(document, "5.3 Electric Service", 2)
    narration_slot(document, "electric_provider")
    heading(document, "5.4 Fire Protection", 2)
    narration_slot(document, "fire_protection")
    heading(document, "5.5 Utility Capacity", 2)
    para(document, "{{ utility_capacity }}")
    heading(document, "5.6 Right-of-Way", 2)
    para(document, "{{ row_info }}")
    heading(document, "5.7 Transportation & Access", 2)
    para(document, "{{ transportation_reqs }}")

    heading(document, "6 Feasibility Determination", 1)
    para(
        document,
        "The recommendations below are ordered by dependency: unresolved confirmations "
        "first, then required studies, optimizations, and permits to initiate.",
    )
    for index in range(1, 6):
        item = document.add_paragraph(style="List Number")
        item.add_run(f"{{{{ recommendation_{index} }}}}")


def build_exhibits(document: DocumentType) -> None:
    heading(document, "EXHIBITS", 1)
    para(
        document,
        "Exhibits are referenced in first-citation order. Uploaded exhibit sheets follow "
        "this list; generated map sheets are appended as they become available.",
    )
    label_value_table(
        document,
        [(f"EXHIBIT {index}", f"{{{{ exhibit_{index} }}}}") for index in range(1, 6)],
    )


def build_footer(document: DocumentType) -> None:
    footer = document.sections[-1].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _meta(text: str) -> None:
        run = paragraph.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = CAPTION
        run.font.name = BODY_FONT

    _meta("{{ project_name }} · {{ report_date }} · Page ")
    _field(paragraph, "PAGE")
    _meta(" of ")
    _field(paragraph, "NUMPAGES")
    _meta(" · Prepared with Civil1")


def build(out_path: Path) -> None:
    document = docx.Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for edge in ("top_margin", "bottom_margin"):
        setattr(section, edge, Inches(0.9))
    for edge in ("left_margin", "right_margin"):
        setattr(section, edge, Inches(1.0))

    define_styles(document)
    build_cover(document)
    build_exec_summary(document)
    build_toc(document)
    build_body(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    build_exhibits(document)
    build_footer(document)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    print(f"wrote {out_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    default_out = Path(__file__).resolve().parents[1] / "assets" / "templates" / "civil1_study_v1.docx"
    parser.add_argument("--out", type=Path, default=default_out)
    args = parser.parse_args()
    build(args.out)


if __name__ == "__main__":
    main()
