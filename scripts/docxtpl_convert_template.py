"""One-time conversion: the firm's ACE feasibility template -> a docxtpl asset.

E1 spike (2026-07-15): the ACE template (`client-data/ATXCivil_Feasibility_Template.docx`)
is a slot-substitution skeleton whose `{TOKEN}` placeholders are Word-native curly-brace
text, not jinja tags. This script rewrites every `{TOKEN}` run in-place to docxtpl's
`{{ token }}` syntax (lowercased, valid Python identifier) so `DocxTemplate.render()` can
fill it, while leaving every other run (fonts, bold, headings, numbering) untouched.

Verified 2026-07-15 against the source template: every `{TOKEN}` sits intact within a
single python-docx run (no run-splitting across a token) -- confirmed by comparing tokens
found in each paragraph's full text against tokens found in its runs' concatenated text,
zero mismatches across all 107 paragraphs + the one table. docxtpl's usual run-splitting
gotcha (a tag broken across runs by Word's spell-check/autocorrect re-flowing) does not
apply here, so a straightforward per-run regex substitution is safe. If a future edit to
the source template (e.g. someone retypes a token in Word) reintroduces run-splitting,
this script's verification step will start reporting unconverted `{TOKEN}` residue and it
needs the docxtpl "fix broken tags" recipe (re-type the tag in Word to normalize runs, or
merge the run text before substituting).

Tokens are converted 1:1 (`{PROPERTY_ADDRESS}` -> `{{ property_address }}`) except the
five image/table slots noted in NOT_CONVERTED below, which need dedicated handling (jinja
loops / InlineImage) in a later pass -- not a run-level text substitution.

Usage::

    uv run python scripts/docxtpl_convert_template.py
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

REPO_ROOT = Path(__file__).resolve().parent.parent
SOURCE_TEMPLATE = (
    REPO_ROOT.parent / "client-data" / "ATXCivil_Feasibility_Template.docx"
)
OUTPUT_TEMPLATE = REPO_ROOT / "assets" / "templates" / "atxcivil_v1.docx"

_TOKEN_RE = re.compile(r"\{([A-Z_0-9]+)\}")

# {TOKEN} instances not handled by the 1:1 substitution below:
# - "![LOGO/IMAGE]" is the cover *site aerial* slot (every delivered study in the corpus
#   has the parcel aerial there) -- rewritten to `{{ cover_aerial }}` and fed an
#   InlineImage by the renderer.
# - The identification table (TRACT/PARCEL ID/ADDRESS/DEED DOC. NO./ACRES) has no
#   tokens in its data row today -- it's populated by hand in every delivered study
#   seen in the corpus. Wiring it needs a jinja loop over tract rows, deferred to E2.
#
# Fidelity pass (2026-07-19): the raw template file ACE shared diverges from their
# *delivered* reports (client-data/feasibility-studies/*.pdf). The delivered look is the
# contract, so after token substitution we reshape the cover to match the corpus:
# banner logo moves off the cover into a small centered running header, cover text is
# centered, cover titles leave the Heading-1 outline (they broke the linter's
# heading-sequence check), and fonts pin to the Calisto MT family their PDFs embed
# (raw template fell back to Arial docDefaults).

BODY_FONT = "Calisto MT"
HEADER_LOGO_WIDTH = Inches(2.0)

# Tokens the render script feeds a docxtpl Subdoc (real multi-paragraph narration,
# not a plain string) -- these need docxtpl's paragraph-tag syntax `{{p token }}`,
# not the inline `{{ token }}` used for every other token. A Subdoc silently
# renders as nothing under an inline tag (verified against the E1 spike output --
# every subdoc-fed section rendered as a blank paragraph until this was fixed).
# Each of these already sits alone in its own paragraph in the source template
# (verified below at conversion time), which paragraph tags require.
SUBDOC_TOKENS = frozenset(
    {
        "ZONING_REGS",
        "WATER_SERVICE",
        "WASTEWATER_SERVICE",
        "ELECTRIC_PROVIDER",
        "FIRE_PROTECTION",
        "FLOODPLAIN_STATUS",
    }
)


def _jinja_name(token: str) -> str:
    return token.lower()


def _delete_paragraph(paragraph: docx.text.paragraph.Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _set_style_font(document: docx.document.Document, style_name: str, font_name: str) -> None:
    from docx.oxml.ns import qn

    style = document.styles[style_name]
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        from docx.oxml import OxmlElement

        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), font_name)


def fidelity_pass(document: docx.document.Document, logo_bytes: bytes | None) -> None:
    """Reshape the converted file to match ACE's *delivered* reports, not the raw
    template skeleton (corpus: client-data/feasibility-studies/*.pdf).

    1. Banner logo: off the cover, into a small centered running header on body pages
       (cover uses a first-page header with no logo, like every delivered study).
    2. Cover aerial: the literal "![LOGO/IMAGE]" line becomes `{{ cover_aerial }}`.
    3. Cover block: centered, and cover titles leave the Heading-1 outline so the
       linter's heading-sequence check sees only the numbered body outline.
    4. Fonts: pin Normal + headings to Calisto MT (embedded in their delivered PDFs;
       the raw template fell back to Arial docDefaults).
    5. Page break after the client block so the cover stands alone.
    """
    if logo_bytes:
        section = document.sections[0]
        section.different_first_page_header_footer = True
        header = section.header
        header.is_linked_to_previous = False
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.add_run().add_picture(BytesIO(logo_bytes), width=HEADER_LOGO_WIDTH)

    # Cover banner paragraph is the first paragraph holding an inline drawing.
    for paragraph in document.paragraphs[:4]:
        if paragraph._p.xpath(".//w:drawing"):
            _delete_paragraph(paragraph)
            break

    body_start = None
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text == "FEASIBILITY STUDY FOR":
            body_start = index
            break
    cover_end = body_start if body_start is not None else 10

    for paragraph in document.paragraphs[:cover_end]:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if (paragraph.style.name or "").startswith("Heading"):
            # Keep the big-title run formatting; just leave the heading outline.
            paragraph.style = document.styles["Title"]
        if paragraph.text.strip() == "![LOGO/IMAGE]":
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = "{{ cover_aerial }}"

    if body_start is not None:
        from docx.enum.text import WD_BREAK

        prior = document.paragraphs[body_start - 1]
        run = prior.add_run()
        run.add_break(WD_BREAK.PAGE)
        for paragraph in document.paragraphs[body_start : body_start + 2]:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if (paragraph.style.name or "").startswith("Heading"):
                paragraph.style = document.styles["Title"]

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            _set_style_font(document, style_name, BODY_FONT)
        except KeyError:
            continue


def convert(source: Path, dest: Path) -> tuple[int, list[str]]:
    """Rewrite every `{TOKEN}` run to `{{ jinja_name }}` (or `{{p jinja_name }}`
    for SUBDOC_TOKENS that sit alone in their own paragraph).

    Returns the number of tokens converted and any `{TOKEN}`-shaped text left over
    after conversion (should be empty -- a non-empty list means a token survived
    outside a single run, i.e. the run-splitting gotcha this repo's template does
    not currently have).
    """
    document = docx.Document(str(source))
    converted = 0

    def _convert_runs(paragraph: docx.text.paragraph.Paragraph) -> None:
        nonlocal converted
        alone = _TOKEN_RE.fullmatch(paragraph.text.strip())
        use_paragraph_tag = bool(alone) and alone.group(1) in SUBDOC_TOKENS
        for run in paragraph.runs:
            if "{" not in run.text:
                continue

            def _sub(m: re.Match[str]) -> str:
                name = _jinja_name(m.group(1))
                return "{{p " + name + " }}" if use_paragraph_tag else "{{ " + name + " }}"

            new_text, n = _TOKEN_RE.subn(_sub, run.text)
            if n:
                run.text = new_text
                converted += n

    for paragraph in document.paragraphs:
        _convert_runs(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _convert_runs(paragraph)

    # The template's one embedded image is the ACE banner logo -- reuse it for the header.
    # (Synthetic test fixtures have no media; the fidelity pass skips the header then.)
    import zipfile

    try:
        with zipfile.ZipFile(source) as archive:
            logo_bytes: bytes | None = archive.read("word/media/image1.png")
    except KeyError:
        logo_bytes = None
    fidelity_pass(document, logo_bytes)

    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dest))

    # Verification pass: re-open the saved file and confirm no {TOKEN} text remains
    # (the image placeholder is expected residue, everything else is not).
    saved = docx.Document(str(dest))
    residue = []
    for paragraph in saved.paragraphs:
        for match in _TOKEN_RE.finditer(paragraph.text):
            if match.group(0) != "{LOGO/IMAGE}":
                residue.append(match.group(0))
    return converted, residue


def main() -> int:
    if not SOURCE_TEMPLATE.exists():
        raise SystemExit(f"Source template not found: {SOURCE_TEMPLATE}")
    converted, residue = convert(SOURCE_TEMPLATE, OUTPUT_TEMPLATE)
    print(f"Converted {converted} token(s) -> {OUTPUT_TEMPLATE}")
    if residue:
        print(f"WARNING: {len(residue)} unconverted token(s) remain: {sorted(set(residue))}")
        return 1
    print("Verification passed: no unconverted {TOKEN} residue.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
