"""E6: export-fidelity linter for exported feasibility study DOCX files.

`python-docx`-based checks that catch the exact failure classes found by hand
comparing real app exports against client-delivered studies (see
`HANDOFF-2026-07-DATA-AND-EXPORT.md`): structural drift from the ACE template
outline, leaked template/placeholder syntax, unit-conversion bugs printed as
plausible-looking text, an internal UUID surfacing as a client-facing ID (D14),
same-document contradictions, and truncated output.

Until E2 (the platform export endpoint) exists, this is the manual review tool
for Brian's app-exported samples -- wire it into CI for the export service once
E2 ships.

Usage::

    uv run python scripts/lint_export_docx.py path/to/exported.docx [more.docx ...]
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

import docx

# The ACE template's outline, in order (client-data/ATXCivil_Feasibility_Template.docx,
# verified 2026-07-15). A real export's heading sequence must match this outline-number
# sequence exactly -- the free-form "polish" LLM step invents a different structure every
# run, which is the root defect E1-E5 replace.
EXPECTED_OUTLINE: tuple[str, ...] = (
    "1",
    "2",
    "2.1",
    "2.2",
    "2.3",
    "3",
    "3.1",
    "3.2",
    "3.3",
    "3.4",
    "3.5",
    "3.6",
    "3.7",
    "3.8",
    "3.9",
    "3.10",
    "3.11",
    "3.12",
    "3.13",
    "3.13.1",
    "3.13.2",
    "3.14",
    "3.15",
    "3.16",
    "3.17",
    "3.18",
    "3.19",
    "4",
    "EXHIBITS",
)

_OUTLINE_NUMBER_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s")
_UUID_RE = re.compile(
    r"\b[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}\b", re.IGNORECASE
)
_PLACEHOLDER_LEAK_RE = re.compile(r"\{\{.*?\}\}|\{[A-Z_0-9]+\}|FIRM LOGO", re.IGNORECASE)
# Austin metro ground elevation runs roughly 400-900 ft; a "N feet" claim under this
# floor is implausible for the region and is the signature of an unconverted meters
# value printed under a feet-labeled unit (D8's bug class).
_AUSTIN_METRO_MIN_PLAUSIBLE_FT = 260
_ELEVATION_FT_RE = re.compile(r"([\d,]+(?:\.\d+)?)\s*(?:ft|feet)\b", re.IGNORECASE)
_CONTRADICTION_PAIRS: tuple[tuple[str, str], ...] = (
    (r"\bunplatted\b", r"\bcould not be confirmed\b"),
    (r"\bplatted\b(?!.*not)", r"\bunplatted\b"),
)


@dataclass(frozen=True)
class Finding:
    check: str
    detail: str


def _full_text(document: docx.document.Document) -> str:
    return "\n".join(p.text for p in document.paragraphs)


def _headings(document: docx.document.Document) -> list[tuple[str, str]]:
    """Returns (outline_number, heading_text) for every Heading-styled paragraph."""
    out: list[tuple[str, str]] = []
    for p in document.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        text = p.text.strip()
        if text.upper() == "EXHIBITS":
            out.append(("EXHIBITS", text))
            continue
        match = _OUTLINE_NUMBER_RE.match(text)
        if match:
            out.append((match.group(1), text))
    return out


def check_heading_sequence(document: docx.document.Document) -> list[Finding]:
    headings = _headings(document)
    found = [number for number, _ in headings]
    if tuple(found) == EXPECTED_OUTLINE:
        return []
    return [
        Finding(
            "heading_sequence",
            f"heading outline does not match the ACE template: expected {list(EXPECTED_OUTLINE)}, "
            f"got {found}",
        )
    ]


def check_placeholder_leaks(document: docx.document.Document) -> list[Finding]:
    text = _full_text(document)
    leaks = sorted(set(_PLACEHOLDER_LEAK_RE.findall(text)))
    return [Finding("placeholder_leak", f"leaked template syntax: {leak!r}") for leak in leaks]


def check_unit_sanity(document: docx.document.Document) -> list[Finding]:
    findings = []
    for p in document.paragraphs:
        if "elevation" not in p.text.lower():
            continue
        for raw in _ELEVATION_FT_RE.findall(p.text):
            value = float(raw.replace(",", ""))
            if value < _AUSTIN_METRO_MIN_PLAUSIBLE_FT:
                findings.append(
                    Finding(
                        "unit_sanity",
                        f"implausible elevation {value:g} ft (Austin metro floor is "
                        f"~{_AUSTIN_METRO_MIN_PLAUSIBLE_FT} ft) -- likely an unconverted "
                        f"meters value: {p.text[:120]!r}",
                    )
                )
    return findings


def check_uuid_leak(document: docx.document.Document) -> list[Finding]:
    text = _full_text(document)
    leaks = sorted(set(_UUID_RE.findall(text)))
    return [
        Finding("uuid_leak", f"internal UUID surfaced in body text: {leak}") for leak in leaks
    ]


def check_contradictions(document: docx.document.Document) -> list[Finding]:
    text = _full_text(document).lower()
    findings = []
    for pattern_a, pattern_b in _CONTRADICTION_PAIRS:
        if re.search(pattern_a, text) and re.search(pattern_b, text):
            findings.append(
                Finding(
                    "contradiction",
                    f"document asserts both {pattern_a!r} and {pattern_b!r} -- likely a "
                    f"cross-section consistency bug",
                )
            )
    return findings


def check_truncation(document: docx.document.Document) -> list[Finding]:
    non_empty = [p for p in document.paragraphs if p.text.strip()]
    if not non_empty:
        return [Finding("truncation", "document has no content")]
    last_heading = None
    for p in document.paragraphs:
        if p.style.name.startswith("Heading") or p.text.strip().upper() == "EXHIBITS":
            last_heading = p.text.strip()
    if last_heading is None:
        return [Finding("truncation", "document has no headings at all")]
    if not (last_heading.upper() == "EXHIBITS" or last_heading.startswith("4")):
        return [
            Finding(
                "truncation",
                f"document does not end in the Summary/Exhibits section "
                f"(last heading seen: {last_heading!r}) -- possible mid-document truncation",
            )
        ]
    last_para = non_empty[-1].text.strip()
    if last_para.endswith(("•", "-", ":", ",")):
        return [
            Finding(
                "truncation",
                f"document ends mid-list ({last_para[-40:]!r}) -- possible token-cap truncation",
            )
        ]
    return []


def check_bold_body_ratio(
    document: docx.document.Document, *, threshold: float = 0.5
) -> list[Finding]:
    """The current html-to-docx path bolds nearly everything; post-docxtpl, most body
    text should carry the template's own Normal style, not manual bold runs."""
    total_chars = 0
    bold_chars = 0
    for p in document.paragraphs:
        if p.style.name.startswith("Heading"):
            continue
        for run in p.runs:
            total_chars += len(run.text)
            if run.bold:
                bold_chars += len(run.text)
    if total_chars == 0:
        return []
    ratio = bold_chars / total_chars
    if ratio > threshold:
        return [
            Finding(
                "bold_body_ratio",
                f"{ratio:.0%} of body text is bold (threshold {threshold:.0%}) -- likely the "
                f"html-to-docx path's blanket-bold styling rather than the template's own styles",
            )
        ]
    return []


CHECKS = (
    check_heading_sequence,
    check_placeholder_leaks,
    check_unit_sanity,
    check_uuid_leak,
    check_contradictions,
    check_truncation,
    check_bold_body_ratio,
)


def lint(path: Path) -> list[Finding]:
    document = docx.Document(str(path))
    findings: list[Finding] = []
    for check in CHECKS:
        findings.extend(check(document))
    return findings


def main(argv: list[str] | None = None) -> int:
    args = argv if argv is not None else sys.argv[1:]
    if not args:
        print("Usage: uv run python scripts/lint_export_docx.py <path.docx> [more.docx ...]")
        return 2

    exit_code = 0
    for arg in args:
        path = Path(arg)
        if not path.exists():
            print(f"{path}: FILE NOT FOUND")
            exit_code = 1
            continue
        findings = lint(path)
        if not findings:
            print(f"{path}: PASSED (0 findings)")
            continue
        exit_code = 1
        print(f"{path}: {len(findings)} finding(s)")
        for f in findings:
            print(f"  - [{f.check}] {f.detail}")
    return exit_code


if __name__ == "__main__":
    raise SystemExit(main())
