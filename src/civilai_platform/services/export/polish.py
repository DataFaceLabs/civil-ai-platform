"""Post-render DOCX polish for export skins (Civil1 / ATX presentation defects).

Strips placeholder-only paragraphs, mid-sentence missing sentinels, empty
recommendation/exhibit table rows, orphan ATX labels, empty leaf headings, and
redundant "local street inferred" copy that otherwise hurt DOCX/PDF readability.
"""

from __future__ import annotations

import re
from io import BytesIO

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

from civilai_platform.services.export.context import _MISSING, _PENDING_PLACEHOLDER

_MISSING_TEXT = _MISSING
_PENDING_TEXT = _PENDING_PLACEHOLDER
_PLACEHOLDER_RE = re.compile(
    re.escape(_MISSING_TEXT) + "|" + re.escape(_PENDING_TEXT),
    re.IGNORECASE,
)
_OUTLINE_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s")

# ATX house template emits these as bare normals ahead of real narration.
_ORPHAN_LABELS = frozenset(
    {
        "zoning",
        "zoning district",
        "overlays",
        "impervious cover regulations",
        "impervious cover limit",
        "watershed classification",
        "compatibility standards",
        "allowed uses",
        "utilities",
        "water",
        "water service",
        "wastewater",
        "wastewater service",
        "electric",
        "electric service",
        "fire protection",
        "fire protection requirements",
        "environment",
        "parcel characteristics",
        "ecoregion",
        "hydrology characteristics",
        "floodplain risk",
        "watershed characteristics",
        "edwards aquifer and water quality",
        "drainage, erosion, and water quality",
        "floodplain maps, delineation, and depiction",
        "waterway setback",
        "erosion hazard zone",
        "key facts",
    }
)

# Civil1 static Contents / headings lose "&" under some DOCX round-trips.
_AMP_HEADING_FIXES: tuple[tuple[re.Pattern[str], str], ...] = (
    (re.compile(r"\bPurpose\s+Scope\b"), "Purpose & Scope"),
    (re.compile(r"\bAdjacent Sites\s+Context\b"), "Adjacent Sites & Context"),
    (re.compile(r"\bEntitlements\s+Administration\b"), "Entitlements & Administration"),
    (re.compile(r"\bZoning\s+Land Use\b"), "Zoning & Land Use"),
    (re.compile(r"\bPlatting\s+Subdivision\b"), "Platting & Subdivision"),
    (
        re.compile(r"\bCompatibility\s+Design Standards\b"),
        "Compatibility & Design Standards",
    ),
    (re.compile(r"\bEasements\s+Setbacks\b"), "Easements & Setbacks"),
    (
        re.compile(r"\bSurveys,\s*Title\s+Other Documents\b"),
        "Surveys, Title & Other Documents",
    ),
    (re.compile(r"\bWatershed\s+Waterways\b"), "Watershed & Waterways"),
    (
        re.compile(r"\bSoils,\s*Elevation\s+Topography\b"),
        "Soils, Elevation & Topography",
    ),
    (
        re.compile(r"\bDrainage Areas\s+Design Criteria\b"),
        "Drainage Areas & Design Criteria",
    ),
    (re.compile(r"\bWater Quality\s+Detention\b"), "Water Quality & Detention"),
    (re.compile(r"\bUtilities,\s*Access\s+Mobility\b"), "Utilities, Access & Mobility"),
    (re.compile(r"\bTransportation\s+Access\b"), "Transportation & Access"),
)


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def _style_name(paragraph: Paragraph) -> str:
    style = paragraph.style
    return (style.name if style is not None else "") or ""


def _is_heading(paragraph: Paragraph) -> bool:
    return _style_name(paragraph).lower().startswith("heading")


def _is_missing_only(text: str) -> bool:
    cleaned = text.strip()
    return cleaned == _MISSING_TEXT or cleaned.casefold() == _PENDING_TEXT.casefold()


def _scrub_placeholders(text: str) -> str:
    """Remove missing/pending sentinels and tidy leftover punctuation."""
    cleaned = _PLACEHOLDER_RE.sub("", text)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\s+\.", ".", cleaned)
    cleaned = re.sub(r"\.\.+", ".", cleaned)
    cleaned = re.sub(r"\s+,", ",", cleaned)
    cleaned = re.sub(r",\s*,", ",", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip(" \t,;:")


def scrub_narration_text(text: str) -> str:
    """Grammar / mobility copy fixes applied to body paragraphs and TOC lines."""
    cleaned = text

    cleaned = re.sub(
        r"\bThe property was ([A-Z][^.]{2,120}?)\.",
        r"The property is located in \1.",
        cleaned,
    )
    # Fused stem after placeholder strip: "the property Surface hydrology…"
    cleaned = re.sub(
        r"\bthe property (?=[A-Z][a-z]+)",
        "the property. ",
        cleaned,
    )

    cleaned = re.sub(
        r"Road class:\s*Local street\s*\(inferred\)",
        "Road class: Local Road",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(
        r"\bLocal street\s*\(inferred\)",
        "Local Road",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(
        r"(?i)\s*The road is inferred to be a local street\.?",
        "",
        cleaned,
    )
    cleaned = re.sub(
        r"(?i)\binferred to be a local street\b",
        "classified as a Local Road",
        cleaned,
    )

    for pattern, replacement in _AMP_HEADING_FIXES:
        cleaned = pattern.sub(replacement, cleaned)

    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\s+\.", ".", cleaned)
    cleaned = re.sub(r"\.\.+", ".", cleaned)
    return cleaned.strip()


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_table_row(table: Table, row_index: int) -> None:
    row = table.rows[row_index]._tr
    row.getparent().remove(row)


def _is_orphan_label(text: str) -> bool:
    compact = re.sub(r"\s+", " ", text).strip()
    if not compact or len(compact) > 60:
        return False
    if "." in compact or ":" in compact:
        return False
    return compact.casefold() in _ORPHAN_LABELS


def _outline_token(heading_text: str) -> str | None:
    stripped = heading_text.strip()
    if stripped.upper() == "EXHIBITS":
        return "EXHIBITS"
    match = _OUTLINE_RE.match(stripped)
    return match.group(1) if match else None


def _is_leaf_heading(heading_text: str) -> bool:
    """Leaf = numbered sub-section (has a dot) or unnumbered display heading."""
    token = _outline_token(heading_text)
    if token is None:
        return True  # e.g. "Key Facts", "Contents" line items are not outline parents
    if token == "EXHIBITS":
        return False
    return "." in token


def _next_content_index(paragraphs: list[Paragraph], start: int) -> int | None:
    for index in range(start, len(paragraphs)):
        text = _paragraph_text(paragraphs[index])
        if text:
            return index
    return None


def polish_export_docx(payload: bytes) -> bytes:
    """Remove presentation-only emptiness and scrub known bad stems from a DOCX."""
    document = docx.Document(BytesIO(payload))

    # 1) Drop table rows whose value cell is empty or the missing sentinel.
    #    Also drop fully empty data rows (all cells blank / sentinel).
    for table in list(document.tables):
        for row_index in range(len(table.rows) - 1, -1, -1):
            cells = table.rows[row_index].cells
            if not cells:
                _remove_table_row(table, row_index)
                continue
            cell_texts = [cell.text.strip() for cell in cells]
            scrubbed_cells = [_scrub_placeholders(value) for value in cell_texts]
            if all(not value or _is_missing_only(raw) for value, raw in zip(scrubbed_cells, cell_texts, strict=False)):
                # Keep a header row when it looks like labels (TRACT / PARCEL ID).
                joined = " ".join(cell_texts).upper()
                if row_index == 0 and any(
                    token in joined for token in ("TRACT", "PARCEL", "ADDRESS", "ACRES")
                ):
                    continue
                if row_index > 0 or not any(cell_texts):
                    _remove_table_row(table, row_index)
                    continue
            if len(cells) >= 2:
                value = cell_texts[-1]
                scrubbed = scrubbed_cells[-1]
                if not scrubbed or _is_missing_only(value):
                    _remove_table_row(table, row_index)
                elif scrubbed != value:
                    cells[-1].text = scrubbed

    # 2) Scrub mid-sentence placeholders + narration defects; drop sentinel-only paras.
    for paragraph in list(document.paragraphs):
        text = _paragraph_text(paragraph)
        if not text:
            continue
        if _is_heading(paragraph):
            fixed = scrub_narration_text(text)
            if fixed != text:
                _set_paragraph_text(paragraph, fixed)
            continue
        if _is_missing_only(text):
            _remove_paragraph(paragraph)
            continue
        scrubbed = scrub_narration_text(_scrub_placeholders(text))
        if not scrubbed:
            _remove_paragraph(paragraph)
        elif scrubbed != text:
            _set_paragraph_text(paragraph, scrubbed)

    # 3) Remove orphan ATX label paragraphs (bare "Zoning" ahead of real copy).
    orphans = [
        paragraph
        for paragraph in document.paragraphs
        if not _is_heading(paragraph)
        and _is_orphan_label(_paragraph_text(paragraph))
    ]
    for paragraph in orphans:
        _remove_paragraph(paragraph)

    # 4) Collapse empty leaf headings (heading immediately followed by heading / EOF).
    # Re-snapshot after orphan removal; walk backwards so removals stay stable.
    while True:
        paragraphs = list(document.paragraphs)
        removed = False
        for index in range(len(paragraphs) - 1, -1, -1):
            paragraph = paragraphs[index]
            if not _is_heading(paragraph):
                continue
            text = _paragraph_text(paragraph)
            if not text or not _is_leaf_heading(text):
                continue
            nxt = _next_content_index(paragraphs, index + 1)
            if nxt is None or _is_heading(paragraphs[nxt]):
                _remove_paragraph(paragraph)
                removed = True
                break
        if not removed:
            break

    # 5) Collapse runs of blank paragraphs (keep at most one spacer).
    previous_blank = False
    for paragraph in list(document.paragraphs):
        text = _paragraph_text(paragraph)
        if _is_heading(paragraph):
            previous_blank = False
            continue
        if text:
            previous_blank = False
            continue
        if previous_blank:
            _remove_paragraph(paragraph)
        else:
            previous_blank = True

    output = BytesIO()
    document.save(output)
    return output.getvalue()
