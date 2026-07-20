"""Render an ExportContext through a Word-native docxtpl skin."""

from __future__ import annotations

import base64
import binascii
import re
from io import BytesIO
from typing import Any

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docxcompose.composer import Composer  # type: ignore[import-untyped]
from docxtpl import DocxTemplate, InlineImage  # type: ignore[import-untyped]

from civilai_platform.models.entities import MapExhibit
from civilai_platform.services import artifacts as artifact_svc
from civilai_platform.services.export.context import ExportContext
from civilai_platform.services.export.polish import polish_export_docx
from civilai_platform.services.export.skins import Skin

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_SENTENCE_RE = re.compile(r"(?<=[.!?])\s+(?=[A-Z(\"'])")

# Mega-paragraph safety net: gold ACE studies use short body paras (~2–3 sentences).
_REFLOW_MIN_CHARS = 350
_REFLOW_MIN_SENTENCES = 3
_REFLOW_SENTENCES_PER_PARA = 2

# Per-skin body font for Subdoc narration (must match template / builder).
_SKIN_BODY_FONT: dict[str, str] = {
    "atxcivil_v1": "Calisto MT",
    "civil1_study_v1": "Source Sans 3",
}
_DEFAULT_BODY_FONT = "Calisto MT"
_BODY_FONT_SIZE = Pt(11)


def _reflow_mega_chunk(chunk: str) -> list[str]:
    """Split a wall-of-text chunk into short paragraphs at sentence boundaries."""
    stripped = chunk.strip()
    if len(stripped) < _REFLOW_MIN_CHARS:
        return [stripped]
    sentences = [s.strip() for s in _SENTENCE_RE.split(stripped) if s.strip()]
    if len(sentences) < _REFLOW_MIN_SENTENCES:
        return [stripped]
    paras: list[str] = []
    for index in range(0, len(sentences), _REFLOW_SENTENCES_PER_PARA):
        group = sentences[index : index + _REFLOW_SENTENCES_PER_PARA]
        paras.append(" ".join(group))
    return paras


def _set_run_font(run: Any, font_name: str) -> None:
    run.font.name = font_name
    run.font.size = _BODY_FONT_SIZE
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), font_name)


def _add_runs_with_bold(paragraph: Any, chunk: str, font_name: str) -> None:
    pos = 0
    for match in _BOLD_RE.finditer(chunk):
        if match.start() > pos:
            run = paragraph.add_run(chunk[pos : match.start()])
            _set_run_font(run, font_name)
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        _set_run_font(bold_run, font_name)
        pos = match.end()
    if pos < len(chunk):
        run = paragraph.add_run(chunk[pos:])
        _set_run_font(run, font_name)


def _narration_subdoc(
    template: DocxTemplate,
    text: str,
    *,
    body_font: str = _DEFAULT_BODY_FONT,
) -> Any:
    """Convert the editor's paragraph/bold subset into real Word paragraphs.

    Splits on blank lines (from TipTap ``<p>`` → ``\\n\\n``). Mega single-``<p>``
    walls are reflowed at sentence boundaries. Runs are pinned to the skin body font
    so Subdocs do not fall back to Arial docDefaults.
    """
    subdoc = template.new_subdoc()
    for chunk in re.split(r"\n\s*\n", text):
        chunk = chunk.strip()
        if not chunk:
            continue
        for piece in _reflow_mega_chunk(chunk):
            paragraph = subdoc.add_paragraph()
            try:
                paragraph.style = "Normal"
            except KeyError:
                pass
            _add_runs_with_bold(paragraph, piece, body_font)
    return subdoc


def _thumbnail_bytes(data_url: str | None) -> bytes | None:
    if not data_url or not data_url.startswith("data:image/") or "," not in data_url:
        return None
    header, encoded = data_url.split(",", 1)
    try:
        return base64.b64decode(encoded) if ";base64" in header else encoded.encode()
    except (ValueError, binascii.Error):
        return None


def _exhibit_image(exhibit: MapExhibit) -> bytes | None:
    mime_type = (exhibit.mime_type or "").lower()
    if exhibit.s3_key and (mime_type.startswith("image/") or not mime_type):
        payload = artifact_svc.download_artifact_bytes(exhibit.s3_key)
        if payload:
            return payload
    # PDF previews are generated client-side and persisted specifically so the report
    # can show the uploaded sheet before server-side PDF rasterization lands in X2.
    return _thumbnail_bytes(exhibit.thumbnail_data_url)


def _append_byo_exhibits(rendered: bytes, exhibits: tuple[MapExhibit, ...]) -> bytes:
    if not exhibits:
        return rendered

    master = docx.Document(BytesIO(rendered))
    composer = Composer(master)
    for index, exhibit in enumerate(exhibits, start=1):
        sheet = docx.Document()
        sheet.add_heading(f"EXHIBIT {index} - {exhibit.label or exhibit.name}", level=1)
        image = _exhibit_image(exhibit)
        if image:
            try:
                sheet.add_picture(BytesIO(image), width=Inches(7.0))
            except (ValueError, OSError):
                sheet.add_paragraph(
                    "The uploaded exhibit could not be embedded; use the original project upload."
                )
        else:
            sheet.add_paragraph(
                "The uploaded exhibit is retained with the project but has no embeddable preview."
            )
        composer.append(sheet)

    output = BytesIO()
    composer.save(output)
    return output.getvalue()


def _customer_logo(template: DocxTemplate, payload: bytes | None) -> Any:
    """A cover InlineImage from the tenant's uploaded logo, height-constrained for an
    elegant lockup. Returns "" when absent or unreadable so the `{{ customer_logo }}`
    token renders empty (the firm name below carries the brand) instead of leaking."""
    if not payload:
        return ""
    try:
        return InlineImage(template, BytesIO(payload), height=Inches(0.8))
    except (ValueError, OSError):
        return ""


def _cover_aerial(template: DocxTemplate, exhibits: tuple[MapExhibit, ...]) -> Any:
    """The cover site image (ACE corpus: every delivered study leads with the parcel
    aerial). Prefers an exhibit labeled like an aerial/vicinity map, falls back to the
    first embeddable exhibit, and renders blank when the project has none."""
    def _score(exhibit: MapExhibit) -> int:
        label = f"{exhibit.label or ''} {exhibit.name}".lower()
        return 0 if ("aerial" in label or "vicinity" in label) else 1

    for exhibit in sorted(exhibits, key=_score):
        image = _exhibit_image(exhibit)
        if image:
            try:
                # Sized so the full ACE cover block (title, preparer, aerial, client
                # contact) still fits one page -- 5.8" pushed contact lines to page 2.
                return InlineImage(template, BytesIO(image), width=Inches(5.0))
            except (ValueError, OSError):
                continue
    return ""


def render_docx(context: ExportContext, skin: Skin) -> bytes:
    if not skin.template_path.exists():
        raise FileNotFoundError(f"export skin template not found: {skin.template_path}")

    template = DocxTemplate(str(skin.template_path))
    render_values: dict[str, Any] = dict(context.template_values)
    render_values["customer_logo"] = _customer_logo(template, context.customer_logo)
    render_values["cover_aerial"] = _cover_aerial(template, context.exhibits)
    body_font = _SKIN_BODY_FONT.get(skin.id, _DEFAULT_BODY_FONT)
    for token in skin.narration_tokens:
        text = (context.narration.get(token) or "").strip()
        if not text:
            text = "Not available from current project data."
        render_values[token] = _narration_subdoc(template, text, body_font=body_font)

    # StrictUndefined would turn optional, intentionally unfilled tenant fields into
    # failures. Instead, ensure every declared template variable has an honest value so
    # raw jinja never leaks into a customer document.
    for variable in template.get_undeclared_template_variables():
        render_values.setdefault(variable, "Not available from current project data.")

    template.render(render_values)
    output = BytesIO()
    template.save(output)
    polished = polish_export_docx(output.getvalue())
    return _append_byo_exhibits(polished, context.exhibits)
