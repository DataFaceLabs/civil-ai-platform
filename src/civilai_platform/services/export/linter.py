"""Production wrapper for export fidelity checks (E6), with per-skin outlines."""

from __future__ import annotations

import re
from io import BytesIO

import docx

from civilai_platform.services.export.skins import Skin

_OUTLINE_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s")
_PLACEHOLDER_RE = re.compile(
    r"\{\{.*?\}\}|\{[A-Z_0-9]+\}|!\[[^\]]*\]|FIRM LOGO", re.IGNORECASE
)
_UUID_RE = re.compile(
    r"\b[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}\b",
    re.IGNORECASE,
)
_ELEVATION_RE = re.compile(r"([\d,]+(?:\.\d+)?)\s*(?:ft|feet)\b", re.IGNORECASE)


def _finding(check: str, detail: str) -> dict[str, str]:
    return {"check": check, "detail": detail}


def _is_ordered_subsequence(got: list[str], expected: list[str]) -> bool:
    """True when every item in ``got`` appears in ``expected`` in the same order."""
    if not got:
        return False
    index = 0
    for token in got:
        try:
            index = expected.index(token, index) + 1
        except ValueError:
            return False
    return True


def lint_docx(payload: bytes, skin: Skin) -> list[dict[str, str]]:
    document = docx.Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    findings: list[dict[str, str]] = []

    outline: list[str] = []
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if not style_name.startswith("Heading"):
            continue
        heading = paragraph.text.strip()
        if heading.upper() == "EXHIBITS":
            outline.append("EXHIBITS")
        elif match := _OUTLINE_RE.match(heading):
            outline.append(match.group(1))
    # Empty leaf headings may be polished away; require an ordered subsequence of
    # the skin outline (got ⊆ expected, same order), not byte-identical equality.
    if skin.outline and not _is_ordered_subsequence(outline, list(skin.outline)):
        findings.append(
            _finding(
                "heading_sequence",
                f"outline does not match skin {skin.id}: "
                f"expected subsequence of {list(skin.outline)}, got {outline}",
            )
        )

    for leak in sorted(set(_PLACEHOLDER_RE.findall(text))):
        findings.append(_finding("placeholder_leak", f"leaked template syntax: {leak!r}"))
    for value in sorted(set(_UUID_RE.findall(text))):
        findings.append(_finding("uuid_leak", f"internal UUID surfaced in body text: {value}"))

    lowered = text.lower()
    if "unplatted" in lowered and "platted" in lowered.replace("unplatted", ""):
        findings.append(_finding("contradiction", "document asserts both platted and unplatted"))

    for paragraph in document.paragraphs:
        if "elevation" not in paragraph.text.lower():
            continue
        for raw in _ELEVATION_RE.findall(paragraph.text):
            value = float(raw.replace(",", ""))
            if value < 260:
                findings.append(
                    _finding("unit_sanity", f"implausible Austin-metro elevation {value:g} ft")
                )

    non_empty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    if not non_empty:
        findings.append(_finding("truncation", "document has no content"))
    elif non_empty[-1].text.strip().endswith(("•", "-", ":", ",")):
        findings.append(_finding("truncation", "document appears to end mid-list"))

    total_chars = 0
    bold_chars = 0
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            continue
        for run in paragraph.runs:
            total_chars += len(run.text)
            if run.bold:
                bold_chars += len(run.text)
    if total_chars and bold_chars / total_chars > 0.5:
        findings.append(
            _finding("bold_body_ratio", f"{bold_chars / total_chars:.0%} of body text is bold")
        )
    return findings
